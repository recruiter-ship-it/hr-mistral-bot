"""
MCP (Model Context Protocol) Client for HR Bot
Полноценная реализация MCP клиента как в OpenClaw

MCP позволяет подключать внешние инструменты (skills) к AI агенту:
- Tools: функции для выполнения задач
- Resources: данные и файлы
- Prompts: шаблоны промптов
"""

import os
import json
import logging
import asyncio
import subprocess
import sys
from typing import Dict, List, Any, Optional, Callable
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from enum import Enum

logger = logging.getLogger(__name__)

# ============================================================
# MCP TYPES
# ============================================================

class MCPTransport(Enum):
    STDIO = "stdio"
    HTTP = "http"
    WEBSOCKET = "websocket"


@dataclass
class MCPTool:
    """Определение инструмента MCP"""
    name: str
    description: str
    input_schema: Dict  # JSON Schema для параметров
    
    def to_mistral_tool(self) -> Dict:
        """Конвертация в формат Mistral"""
        return {
            "type": "function",
            "function": {
                "name": self.name,
                "description": self.description,
                "parameters": self.input_schema
            }
        }


@dataclass
class MCPResource:
    """Ресурс MCP (файл, данные)"""
    uri: str
    name: str
    description: str = ""
    mime_type: str = "text/plain"


@dataclass
class MCPPrompt:
    """Шаблон промпта MCP"""
    name: str
    description: str
    arguments: List[Dict] = field(default_factory=list)


@dataclass
class MCPServerConfig:
    """Конфигурация MCP сервера"""
    name: str
    description: str = ""
    command: str = ""  # Для stdio
    args: List[str] = field(default_factory=list)
    url: str = ""  # Для HTTP/WebSocket
    transport: MCPTransport = MCPTransport.STDIO
    env: Dict[str, str] = field(default_factory=dict)
    enabled: bool = True
    
    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "description": self.description,
            "command": self.command,
            "args": self.args,
            "url": self.url,
            "transport": self.transport.value,
            "env": self.env,
            "enabled": self.enabled
        }
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'MCPServerConfig':
        # Подставляем реальные значения env переменных
        env = data.get("env", {})
        resolved_env = {}
        for key, value in env.items():
            if isinstance(value, str) and value.startswith("${") and value.endswith("}"):
                # Извлекаем имя переменной из ${VAR_NAME}
                var_name = value[2:-1]
                resolved_env[key] = os.environ.get(var_name, "")
            else:
                resolved_env[key] = value
        
        return cls(
            name=data["name"],
            description=data.get("description", ""),
            command=data.get("command", ""),
            args=data.get("args", []),
            url=data.get("url", ""),
            transport=MCPTransport(data.get("transport", "stdio")),
            env=resolved_env,
            enabled=data.get("enabled", True)
        )


# ============================================================
# MCP SERVER CONNECTION
# ============================================================

class MCPServerConnection:
    """Соединение с MCP сервером"""
    
    def __init__(self, config: MCPServerConfig):
        self.config = config
        self.process: Optional[subprocess.Popen] = None
        self.tools: List[MCPTool] = []
        self.resources: List[MCPResource] = []
        self.prompts: List[MCPPrompt] = []
        self.connected = False
        self._request_id = 0
    
    async def connect(self) -> bool:
        """Подключение к MCP серверу"""
        if self.config.transport == MCPTransport.STDIO:
            return await self._connect_stdio()
        elif self.config.transport == MCPTransport.HTTP:
            return await self._connect_http()
        return False
    
    async def _connect_stdio(self) -> bool:
        """Подключение через stdio (запуск процесса)"""
        try:
            # Подготавливаем окружение
            env = os.environ.copy()
            env.update(self.config.env)
            
            # Запускаем процесс
            cmd = [self.config.command] + self.config.args
            
            # Если команда начинается с npx или python, используем shell
            if self.config.command in ["npx", "npm", "uvx"]:
                self.process = subprocess.Popen(
                    " ".join(cmd),
                    stdin=subprocess.PIPE,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    shell=True,
                    env=env
                )
            else:
                self.process = subprocess.Popen(
                    cmd,
                    stdin=subprocess.PIPE,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    env=env
                )
            
            # Инициализация MCP
            await self._send_request("initialize", {
                "protocolVersion": "2024-11-05",
                "clientInfo": {
                    "name": "hr-bot",
                    "version": "1.0.0"
                },
                "capabilities": {
                    "tools": {},
                    "resources": {},
                    "prompts": {}
                }
            })
            
            # Получаем список инструментов
            await self._load_tools()
            
            self.connected = True
            logger.info(f"Connected to MCP server: {self.config.name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to connect to MCP server {self.config.name}: {e}")
            return False
    
    async def _connect_http(self) -> bool:
        """Подключение через HTTP"""
        try:
            import aiohttp
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.config.url}/initialize",
                    json={
                        "protocolVersion": "2024-11-05",
                        "clientInfo": {"name": "hr-bot", "version": "1.0.0"}
                    }
                ) as response:
                    if response.status == 200:
                        await self._load_tools_http(session)
                        self.connected = True
                        return True
        except Exception as e:
            logger.error(f"HTTP connection failed: {e}")
        return False
    
    async def _load_tools_http(self, session):
        """Загрузка инструментов через HTTP"""
        async with session.get(f"{self.config.url}/tools") as response:
            if response.status == 200:
                data = await response.json()
                for tool in data.get("tools", []):
                    self.tools.append(MCPTool(
                        name=tool["name"],
                        description=tool.get("description", ""),
                        input_schema=tool.get("inputSchema", {})
                    ))
    
    async def _send_request(self, method: str, params: Dict = None) -> Optional[Dict]:
        """Отправка JSON-RPC запроса"""
        if not self.process:
            return None
        
        self._request_id += 1
        request = {
            "jsonrpc": "2.0",
            "id": self._request_id,
            "method": method,
            "params": params or {}
        }
        
        try:
            # Отправляем запрос
            request_str = json.dumps(request) + "\n"
            self.process.stdin.write(request_str.encode())
            self.process.stdin.flush()
            
            # Читаем ответ
            response_str = self.process.stdout.readline().decode()
            response = json.loads(response_str)
            
            return response.get("result")
        except Exception as e:
            logger.error(f"MCP request failed: {e}")
            return None
    
    async def _load_tools(self):
        """Загрузка списка инструментов"""
        result = await self._send_request("tools/list")
        if result:
            for tool in result.get("tools", []):
                self.tools.append(MCPTool(
                    name=tool["name"],
                    description=tool.get("description", ""),
                    input_schema=tool.get("inputSchema", {})
                ))
            logger.info(f"Loaded {len(self.tools)} tools from {self.config.name}")
    
    async def call_tool(self, tool_name: str, arguments: Dict) -> Any:
        """Вызов инструмента"""
        if not self.connected:
            return {"error": "Not connected to server"}
        
        if self.config.transport == MCPTransport.STDIO:
            result = await self._send_request("tools/call", {
                "name": tool_name,
                "arguments": arguments
            })
            
            if result:
                # Извлекаем контент из результата
                content = result.get("content", [])
                if content:
                    # Объединяем все текстовые результаты
                    texts = []
                    for item in content:
                        if item.get("type") == "text":
                            texts.append(item.get("text", ""))
                    return "\n".join(texts)
            return result
        
        elif self.config.transport == MCPTransport.HTTP:
            return await self._call_tool_http(tool_name, arguments)
        
        return {"error": "Unknown transport"}
    
    async def _call_tool_http(self, tool_name: str, arguments: Dict) -> Any:
        """Вызов инструмента через HTTP"""
        try:
            import aiohttp
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{self.config.url}/tools/{tool_name}/call",
                    json={"arguments": arguments}
                ) as response:
                    if response.status == 200:
                        return await response.json()
        except Exception as e:
            return {"error": str(e)}
    
    async def disconnect(self):
        """Отключение от сервера"""
        if self.process:
            self.process.terminate()
            self.process = None
        self.connected = False


# ============================================================
# MCP CLIENT MANAGER
# ============================================================

class MCPClientManager:
    """
    Менеджер MCP клиентов - управляет подключениями к серверам
    Аналог системы Skills в OpenClaw
    """
    
    def __init__(self, config_path: str = None):
        self.config_path = config_path or "mcp_config.json"
        self.servers: Dict[str, MCPServerConnection] = {}
        self.tool_to_server: Dict[str, str] = {}  # tool_name -> server_name
        self._load_config()
    
    def _load_config(self):
        """Загрузка конфигурации MCP серверов"""
        config_file = Path(self.config_path)
        if config_file.exists():
            try:
                with open(config_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for server_data in data.get("mcpServers", []):
                        config = MCPServerConfig.from_dict(server_data)
                        self.servers[config.name] = MCPServerConnection(config)
                logger.info(f"Loaded {len(self.servers)} MCP server configs")
            except Exception as e:
                logger.error(f"Failed to load MCP config: {e}")
    
    def _save_config(self):
        """Сохранение конфигурации"""
        data = {
            "mcpServers": [s.config.to_dict() for s in self.servers.values()]
        }
        with open(self.config_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
    
    async def add_server(self, config: MCPServerConfig) -> bool:
        """Добавление MCP сервера"""
        if config.name in self.servers:
            logger.warning(f"Server {config.name} already exists")
            return False
        
        connection = MCPServerConnection(config)
        if await connection.connect():
            self.servers[config.name] = connection
            # Маппинг инструментов
            for tool in connection.tools:
                self.tool_to_server[tool.name] = config.name
            self._save_config()
            return True
        return False
    
    def remove_server(self, name: str) -> bool:
        """Удаление сервера"""
        if name in self.servers:
            conn = self.servers[name]
            asyncio.create_task(conn.disconnect())
            # Удаляем маппинг инструментов
            for tool in conn.tools:
                self.tool_to_server.pop(tool.name, None)
            del self.servers[name]
            self._save_config()
            return True
        return False
    
    async def connect_all(self) -> Dict[str, bool]:
        """Подключение ко всем серверам"""
        results = {}
        for name, connection in self.servers.items():
            if connection.config.enabled:
                success = await connection.connect()
                results[name] = success
                if success:
                    for tool in connection.tools:
                        self.tool_to_server[tool.name] = name
        return results
    
    def get_all_tools(self) -> List[Dict]:
        """Получение всех инструментов для Mistral"""
        tools = []
        for connection in self.servers.values():
            if connection.connected:
                for tool in connection.tools:
                    tools.append(tool.to_mistral_tool())
        return tools
    
    def get_tool_names(self) -> List[str]:
        """Список всех инструментов"""
        return list(self.tool_to_server.keys())
    
    async def call_tool(self, tool_name: str, arguments: Dict) -> Any:
        """Вызов инструмента по имени"""
        server_name = self.tool_to_server.get(tool_name)
        if not server_name:
            return {"error": f"Tool {tool_name} not found"}
        
        connection = self.servers.get(server_name)
        if not connection or not connection.connected:
            return {"error": f"Server {server_name} not connected"}
        
        return await connection.call_tool(tool_name, arguments)
    
    def list_servers(self) -> List[Dict]:
        """Список серверов"""
        return [
            {
                "name": name,
                "connected": conn.connected,
                "tools_count": len(conn.tools),
                "enabled": conn.config.enabled
            }
            for name, conn in self.servers.items()
        ]


# ============================================================
# BUILT-IN MCP SERVERS (Local Implementation)
# ============================================================

class LocalMCPServer:
    """
    Локальный MCP-подобный сервер (без отдельного процесса)
    Для встроенных навыков
    """
    
    def __init__(self, name: str, description: str):
        self.name = name
        self.description = description
        self.tools: Dict[str, Callable] = {}
        self.tool_schemas: Dict[str, Dict] = {}
    
    def register_tool(self, name: str, handler: Callable, schema: Dict):
        """Регистрация инструмента"""
        self.tools[name] = handler
        self.tool_schemas[name] = schema
    
    def get_tools(self) -> List[MCPTool]:
        """Получение списка инструментов"""
        return [
            MCPTool(
                name=name,
                description=schema.get("description", ""),
                input_schema=schema.get("parameters", {})
            )
            for name, schema in self.tool_schemas.items()
        ]
    
    async def call_tool(self, name: str, arguments: Dict) -> Any:
        """Вызов инструмента"""
        handler = self.tools.get(name)
        if not handler:
            return {"error": f"Tool {name} not found"}
        
        try:
            if asyncio.iscoroutinefunction(handler):
                result = await handler(**arguments)
            else:
                result = handler(**arguments)
            return result
        except Exception as e:
            return {"error": str(e)}


# ============================================================
# HR MCP SERVERS
# ============================================================

def create_documents_mcp_server() -> LocalMCPServer:
    """Создание MCP сервера для работы с документами"""
    server = LocalMCPServer("documents", "Работа с документами Office")
    
    def create_document(title: str, content: str, doc_type: str = "docx") -> Dict:
        """Создание документа"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            import io
            
            doc = Document()
            
            # Добавляем заголовок
            doc.add_heading(title, level=1)
            
            # Добавляем содержимое
            for line in content.split('\n'):
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
            
            # Сохраняем
            filename = f"{title.replace(' ', '_')}.{doc_type}"
            filepath = Path("skills/documents") / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(filepath))
            
            return {
                "success": True,
                "filename": filename,
                "filepath": str(filepath),
                "message": f"✅ Документ создан: {filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def create_spreadsheet(title: str, data: List[List], filename: str = None) -> Dict:
        """Создание таблицы"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            
            wb = Workbook()
            ws = wb.active
            ws.title = title[:31]
            
            for row_idx, row in enumerate(data, 1):
                for col_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    if row_idx == 1:
                        cell.font = Font(bold=True)
            
            if not filename:
                filename = f"{title.replace(' ', '_')}.xlsx"
            
            filepath = Path("skills/documents") / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(filepath))
            
            return {
                "success": True,
                "filename": filename,
                "filepath": str(filepath),
                "message": f"✅ Таблица создана: {filename}"
            }
        except ImportError:
            return {"success": False, "error": "openpyxl not installed"}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def read_document(filepath: str) -> Dict:
        """Чтение документа"""
        try:
            from docx import Document
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return {"success": True, "content": text}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    server.register_tool("create_document", create_document, {
        "description": "Создать документ Word с заголовком и содержимым",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Заголовок документа"},
                "content": {"type": "string", "description": "Содержимое (поддерживает Markdown)"},
                "doc_type": {"type": "string", "description": "Тип документа (docx, pdf)", "default": "docx"}
            },
            "required": ["title", "content"]
        }
    })
    
    server.register_tool("create_spreadsheet", create_spreadsheet, {
        "description": "Создать таблицу Excel с данными",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Название листа"},
                "data": {"type": "array", "items": {"type": "array"}, "description": "Данные таблицы"},
                "filename": {"type": "string", "description": "Имя файла"}
            },
            "required": ["title", "data"]
        }
    })
    
    server.register_tool("read_document", read_document, {
        "description": "Прочитать содержимое документа",
        "parameters": {
            "type": "object",
            "properties": {
                "filepath": {"type": "string", "description": "Путь к файлу"}
            },
            "required": ["filepath"]
        }
    })
    
    return server


def create_hr_mcp_server() -> LocalMCPServer:
    """Создание MCP сервера для HR задач"""
    server = LocalMCPServer("hr", "HR инструменты: офферы, welcome-письма, кандидатам")
    
    def create_offer(candidate_name: str, position: str, salary: str, 
                     start_date: str, department: str = "", company: str = "Компания") -> Dict:
        """Создание оффера"""
        from datetime import datetime
        
        content = f"""# ОФФЕР О ПРИНЯТИИ НА РАБОТУ

**Компания:** {company}  
**Дата:** {datetime.now().strftime("%d.%m.%Y")}

---

## Уважаемый(ая) {candidate_name}!

Мы рады предложить Вам должность **{position}** в команде {company}.

### Условия предложения:

| Параметр | Значение |
|----------|----------|
| **Должность** | {position} |
| **Отдел** | {department or 'Не указан'} |
| **Тип занятости** | Полная занятость |
| **Дата выхода** | {start_date} |
| **Испытательный срок** | 3 месяца |
| **Зарплата** | {salary} |

---

Для принятия предложения, пожалуйста, подтвердите в течение 3 рабочих дней.

**С уважением,**  
HR Team  
{company}
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Offer_{candidate_name.replace(' ', '_')}.md"
        }
    
    def create_welcome(employee_name: str, position: str, start_date: str,
                       start_time: str = "10:00", buddy: str = "", 
                       manager: str = "", company: str = "Компания") -> Dict:
        """Создание welcome-письма"""
        content = f"""# Добро пожаловать в команду! 🎉

**Привет, {employee_name}!**

Поздравляем с присоединением к команде {company}!

---

## 📅 Твой первый день

**Дата:** {start_date}  
**Время:** {start_time}  

---

## 📋 Что взять с собой:
- Паспорт
- ИНН
- СНИЛС
- Диплом
- Трудовую книжку

---

## 🗓️ План первой недели:

**День 1:** Знакомство с командой, оформление  
**День 2:** Обучение продуктам  
**День 3:** Знакомство с отделами  
**День 4:** Обучение инструментам  
**День 5:** Первые задачи, 1-on-1 с руководителем

---
"""
        if buddy:
            content += f"- **Buddy:** {buddy}\n"
        if manager:
            content += f"- **Руководитель:** {manager}\n"
        
        return {
            "success": True,
            "content": content,
            "filename": f"Welcome_{employee_name.replace(' ', '_')}.md"
        }
    
    def create_rejection(candidate_name: str, position: str, 
                         keep_in_touch: bool = True, company: str = "Компания") -> Dict:
        """Создание письма с отказом"""
        keep_text = "\n- Сохранить Ваше резюме в базе\n" if keep_in_touch else ""
        
        content = f"""# Уважаемый(ая) {candidate_name}!

Благодарим за интерес к вакансии **{position}** в {company}.

К сожалению, на данном этапе мы не можем предложить Вам эту должность.

Это не означает, что Ваш опыт не представляет ценности — мы ищем кандидата с другим профилем.
{keep_text}
Желаем успехов в поиске!

С уважением,  
HR Team
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Rejection_{candidate_name.replace(' ', '_')}.md"
        }
    
    def create_interview_invite(candidate_name: str, position: str,
                                interview_date: str, interview_time: str,
                                duration: int = 60, interview_type: str = "онлайн",
                                company: str = "Компания") -> Dict:
        """Создание приглашения на интервью"""
        content = f"""# Приглашение на интервью

**Уважаемый(ая) {candidate_name}!**

Благодарим за интерес к вакансии **{position}** в {company}.

## 📅 Детали интервью:

| Параметр | Значение |
|----------|----------|
| **Дата** | {interview_date} |
| **Время** | {interview_time} |
| **Формат** | {interview_type} |
| **Длительность** | {duration} минут |

---

Пожалуйста, подтвердите участие ответным письмом.

С уважением,  
HR Team
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Interview_{candidate_name.replace(' ', '_')}.md"
        }
    
    # Регистрируем инструменты
    server.register_tool("create_offer", create_offer, {
        "description": "Создать оффер о приёме на работу",
        "parameters": {
            "type": "object",
            "properties": {
                "candidate_name": {"type": "string", "description": "Имя кандидата"},
                "position": {"type": "string", "description": "Должность"},
                "salary": {"type": "string", "description": "Зарплата"},
                "start_date": {"type": "string", "description": "Дата выхода"},
                "department": {"type": "string", "description": "Отдел"},
                "company": {"type": "string", "description": "Компания"}
            },
            "required": ["candidate_name", "position", "salary", "start_date"]
        }
    })
    
    server.register_tool("create_welcome_letter", create_welcome, {
        "description": "Создать welcome-письмо для нового сотрудника",
        "parameters": {
            "type": "object",
            "properties": {
                "employee_name": {"type": "string", "description": "Имя сотрудника"},
                "position": {"type": "string", "description": "Должность"},
                "start_date": {"type": "string", "description": "Дата выхода"},
                "start_time": {"type": "string", "description": "Время выхода"},
                "buddy": {"type": "string", "description": "Имя buddy"},
                "manager": {"type": "string", "description": "Имя руководителя"},
                "company": {"type": "string", "description": "Компания"}
            },
            "required": ["employee_name", "position", "start_date"]
        }
    })
    
    server.register_tool("create_rejection_letter", create_rejection, {
        "description": "Создать письмо с отказом кандидату",
        "parameters": {
            "type": "object",
            "properties": {
                "candidate_name": {"type": "string", "description": "Имя кандидата"},
                "position": {"type": "string", "description": "Должность"},
                "keep_in_touch": {"type": "boolean", "description": "Продолжить общение"},
                "company": {"type": "string", "description": "Компания"}
            },
            "required": ["candidate_name", "position"]
        }
    })
    
    server.register_tool("create_interview_invite", create_interview_invite, {
        "description": "Создать приглашение на интервью",
        "parameters": {
            "type": "object",
            "properties": {
                "candidate_name": {"type": "string", "description": "Имя кандидата"},
                "position": {"type": "string", "description": "Должность"},
                "interview_date": {"type": "string", "description": "Дата интервью"},
                "interview_time": {"type": "string", "description": "Время интервью"},
                "duration": {"type": "integer", "description": "Длительность в минутах"},
                "interview_type": {"type": "string", "description": "Тип (онлайн/офис)"},
                "company": {"type": "string", "description": "Компания"}
            },
            "required": ["candidate_name", "position", "interview_date", "interview_time"]
        }
    })
    
    return server


def create_google_mcp_server() -> LocalMCPServer:
    """Создание MCP сервера для Google Workspace"""
    server = LocalMCPServer("google", "Google Workspace: Sheets, Docs, Calendar")
    
    def add_to_sheet(spreadsheet_id: str, range_name: str, values: List[List]) -> Dict:
        """Добавление данных в Google Sheet"""
        import google_sheets
        
        # Получаем сервис
        service = google_sheets.get_sheets_service()
        if not service:
            return {"success": False, "error": "Google Sheets не настроен"}
        
        try:
            body = {"values": values}
            result = service.spreadsheets().values().append(
                spreadsheetId=spreadsheet_id,
                range=range_name,
                valueInputOption='USER_ENTERED',
                insertDataOption='INSERT_ROWS',
                body=body
            ).execute()
            
            return {
                "success": True,
                "updated_rows": result.get('updates', {}).get('updatedRows', 0),
                "message": f"✅ Добавлено строк: {result.get('updates', {}).get('updatedRows', 0)}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def read_sheet(spreadsheet_id: str, range_name: str) -> Dict:
        """Чтение данных из Google Sheet"""
        import google_sheets
        
        success, data = google_sheets.get_sheet_data(range_name)
        if success:
            return {"success": True, "data": data}
        return {"success": False, "error": data}
    
    def create_google_doc(title: str, content: str) -> Dict:
        """Создание Google Doc"""
        from document_generator import google_docs
        
        result = google_docs.create_document(title, content)
        return result
    
    server.register_tool("add_to_sheet", add_to_sheet, {
        "description": "Добавить данные в Google Sheet",
        "parameters": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string", "description": "ID таблицы"},
                "range_name": {"type": "string", "description": "Диапазон (например, 'Лист1!A:K')"},
                "values": {"type": "array", "items": {"type": "array"}, "description": "Данные для добавления"}
            },
            "required": ["spreadsheet_id", "range_name", "values"]
        }
    })
    
    server.register_tool("read_sheet", read_sheet, {
        "description": "Прочитать данные из Google Sheet",
        "parameters": {
            "type": "object",
            "properties": {
                "spreadsheet_id": {"type": "string", "description": "ID таблицы"},
                "range_name": {"type": "string", "description": "Диапазон"}
            },
            "required": ["spreadsheet_id", "range_name"]
        }
    })
    
    server.register_tool("create_google_doc", create_google_doc, {
        "description": "Создать документ в Google Docs",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Название документа"},
                "content": {"type": "string", "description": "Содержимое (Markdown)"}
            },
            "required": ["title", "content"]
        }
    })
    
    return server


def create_web_mcp_server() -> LocalMCPServer:
    """Создание MCP сервера для веб-запросов"""
    server = LocalMCPServer("web", "Веб-запросы и поиск")
    
    def fetch_url(url: str) -> Dict:
        """Получение содержимого URL"""
        import requests
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return {
                "success": True,
                "content": response.text[:5000],
                "status_code": response.status_code
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def web_search(query: str, num_results: int = 5) -> Dict:
        """Веб-поиск"""
        # Используем Mistral web search через API
        return {
            "success": True,
            "note": "Web search доступен через встроенный web_search tool Mistral",
            "query": query
        }
    
    server.register_tool("fetch_url", fetch_url, {
        "description": "Получить содержимое веб-страницы",
        "parameters": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL страницы"}
            },
            "required": ["url"]
        }
    })
    
    server.register_tool("web_search", web_search, {
        "description": "Поиск в интернете",
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Поисковый запрос"},
                "num_results": {"type": "integer", "description": "Количество результатов"}
            },
            "required": ["query"]
        }
    })
    
    return server


# ============================================================
# MCP ORCHESTRATOR
# ============================================================

class MCPOrchestrator:
    """
    Оркестратор MCP - объединяет внешние и локальные серверы
    Это главный класс для управления всеми навыками агента
    
    Навыки (как в OpenClaw):
    - filesystem: работа с файлами
    - terminal: выполнение команд
    - browser: веб-автоматизация
    - memory: персистентная память
    - communication: Slack, Discord, Email
    - image: генерация изображений
    - database: SQL операции
    - analytics: отчёты и графики
    - documents: создание документов
    - hr: HR документы
    - google: Google Workspace
    - web: веб-запросы
    """
    
    def __init__(self):
        self.client_manager = MCPClientManager()
        self.local_servers: Dict[str, LocalMCPServer] = {}
        self.extended_skills = None  # Новые расширенные навыки
        self.tool_to_server: Dict[str, tuple] = {}  # tool_name -> (server_name, is_local, is_extended)
        
        # Инициализируем локальные серверы
        self._init_local_servers()
        
        # Инициализируем расширенные навыки (как в OpenClaw)
        self._init_extended_skills()
    
    def _init_local_servers(self):
        """Инициализация встроенных MCP серверов"""
        # Регистрируем локальные серверы
        servers = [
            ("documents", create_documents_mcp_server()),
            ("hr", create_hr_mcp_server()),
            ("google", create_google_mcp_server()),
            ("web", create_web_mcp_server()),
        ]
        
        for name, server in servers:
            self.local_servers[name] = server
            # Маппинг инструментов
            for tool in server.get_tools():
                self.tool_to_server[tool.name] = (name, True, False)
        
        logger.info(f"Initialized {len(self.local_servers)} local MCP servers")
    
    def _init_extended_skills(self):
        """Инициализация расширенных навыков (как в OpenClaw)"""
        try:
            from skills_extended import skills_registry
            
            self.extended_skills = skills_registry
            
            # Добавляем инструменты в маппинг
            for skill_name, skill in skills_registry.skills.items():
                for tool in skill.tools:
                    self.tool_to_server[tool.name] = (skill_name, False, True)
            
            logger.info(f"Initialized {len(skills_registry.skills)} extended skills with {len(skills_registry.get_all_tools())} tools")
        except ImportError as e:
            logger.warning(f"Extended skills not available: {e}")
            self.extended_skills = None
    
    async def initialize(self):
        """Полная инициализация"""
        # Подключаем внешние серверы
        results = await self.client_manager.connect_all()
        logger.info(f"External servers connection: {results}")
        
        # Добавляем инструменты внешних серверов в маппинг
        for server_name, connection in self.client_manager.servers.items():
            for tool in connection.tools:
                self.tool_to_server[tool.name] = (server_name, False, False)
    
    def get_all_tools(self) -> List[Dict]:
        """Получение всех инструментов для Mistral"""
        tools = []
        
        # Локальные серверы
        for server in self.local_servers.values():
            for tool in server.get_tools():
                tools.append(tool.to_mistral_tool())
        
        # Расширенные навыки (как в OpenClaw)
        if self.extended_skills:
            tools.extend(self.extended_skills.get_all_tools())
        
        # Внешние серверы
        tools.extend(self.client_manager.get_all_tools())
        
        return tools
    
    def get_tool_names(self) -> List[str]:
        """Список всех инструментов"""
        return list(self.tool_to_server.keys())
    
    async def call_tool(self, tool_name: str, arguments: Dict) -> Any:
        """Вызов инструмента"""
        if tool_name not in self.tool_to_server:
            return {"error": f"Tool {tool_name} not found"}
        
        server_name, is_local, is_extended = self.tool_to_server[tool_name]
        
        if is_extended:
            # Расширенный навык (как в OpenClaw)
            if self.extended_skills:
                return await self.extended_skills.execute_tool(tool_name, **arguments)
        elif is_local:
            # Локальный сервер
            server = self.local_servers.get(server_name)
            if server:
                return await server.call_tool(tool_name, arguments)
        else:
            # Внешний сервер
            return await self.client_manager.call_tool(tool_name, arguments)
        
        return {"error": f"Server {server_name} not found"}
    
    def list_skills(self) -> List[Dict]:
        """Список всех навыков (серверов)"""
        skills = []
        
        # Локальные
        for name, server in self.local_servers.items():
            skills.append({
                "name": name,
                "description": server.description,
                "type": "local",
                "tools_count": len(server.tools),
                "enabled": True
            })
        
        # Расширенные навыки (как в OpenClaw)
        if self.extended_skills:
            for skill_info in self.extended_skills.list_skills():
                skills.append({
                    "name": skill_info["name"],
                    "description": skill_info["description"],
                    "type": "extended",
                    "tools_count": skill_info["tools_count"],
                    "tools": skill_info["tools"],
                    "enabled": True
                })
        
        # Внешние
        for name, conn in self.client_manager.servers.items():
            skills.append({
                "name": name,
                "description": f"External MCP server",
                "type": "external",
                "tools_count": len(conn.tools),
                "enabled": conn.config.enabled,
                "connected": conn.connected
            })
        
        return skills
    
    async def add_external_server(self, config: MCPServerConfig) -> bool:
        """Добавление внешнего MCP сервера"""
        return await self.client_manager.add_server(config)
    
    def remove_external_server(self, name: str) -> bool:
        """Удаление внешнего сервера"""
        return self.client_manager.remove_server(name)


# Глобальный экземпляр
mcp_orchestrator = MCPOrchestrator()
