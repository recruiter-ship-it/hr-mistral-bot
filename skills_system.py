"""
Skills System for HR Bot - MCP-подобная система навыков
Позволяет подключать внешние инструменты и расширять возможности агента
"""

import os
import json
import logging
import importlib
import subprocess
import asyncio
from typing import Dict, List, Any, Optional, Callable
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

logger = logging.getLogger(__name__)

# Директория для навыков
SKILLS_DIR = Path(__file__).parent / "skills"


@dataclass
class Skill:
    """Определение навыка"""
    name: str
    description: str
    version: str = "1.0.0"
    author: str = "unknown"
    tools: List[Dict] = field(default_factory=list)
    dependencies: List[str] = field(default_factory=list)
    config: Dict = field(default_factory=dict)
    enabled: bool = True
    installed_at: datetime = field(default_factory=datetime.now)
    
    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "description": self.description,
            "version": self.version,
            "author": self.author,
            "tools": self.tools,
            "dependencies": self.dependencies,
            "config": self.config,
            "enabled": self.enabled,
            "installed_at": self.installed_at.isoformat()
        }
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'Skill':
        return cls(
            name=data["name"],
            description=data["description"],
            version=data.get("version", "1.0.0"),
            author=data.get("author", "unknown"),
            tools=data.get("tools", []),
            dependencies=data.get("dependencies", []),
            config=data.get("config", {}),
            enabled=data.get("enabled", True),
            installed_at=datetime.fromisoformat(data["installed_at"]) if "installed_at" in data else datetime.now()
        )


class SkillRegistry:
    """Реестр навыков"""
    
    def __init__(self, skills_dir: Path = SKILLS_DIR):
        self.skills_dir = skills_dir
        self.skills: Dict[str, Skill] = {}
        self.tool_handlers: Dict[str, Callable] = {}
        self.config_file = skills_dir / "skills_config.json"
        
        # Создаём директорию если нет
        self.skills_dir.mkdir(parents=True, exist_ok=True)
        
        # Загружаем сохранённые навыки
        self._load_config()
    
    def _load_config(self):
        """Загрузка конфигурации навыков"""
        if self.config_file.exists():
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    for name, skill_data in data.get("skills", {}).items():
                        self.skills[name] = Skill.from_dict(skill_data)
                logger.info(f"Loaded {len(self.skills)} skills from config")
            except Exception as e:
                logger.error(f"Failed to load skills config: {e}")
    
    def _save_config(self):
        """Сохранение конфигурации навыков"""
        try:
            data = {
                "skills": {name: skill.to_dict() for name, skill in self.skills.items()},
                "updated_at": datetime.now().isoformat()
            }
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Failed to save skills config: {e}")
    
    def register_skill(self, skill: Skill, tool_handlers: Dict[str, Callable] = None):
        """Регистрация навыка"""
        self.skills[skill.name] = skill
        
        if tool_handlers:
            for tool_name, handler in tool_handlers.items():
                self.tool_handlers[tool_name] = handler
        
        self._save_config()
        logger.info(f"Registered skill: {skill.name} with {len(skill.tools)} tools")
    
    def unregister_skill(self, name: str) -> bool:
        """Удаление навыка"""
        if name in self.skills:
            skill = self.skills[name]
            # Удаляем обработчики инструментов
            for tool in skill.tools:
                tool_name = tool.get("name")
                if tool_name in self.tool_handlers:
                    del self.tool_handlers[tool_name]
            
            del self.skills[name]
            self._save_config()
            logger.info(f"Unregistered skill: {name}")
            return True
        return False
    
    def get_tool_handler(self, tool_name: str) -> Optional[Callable]:
        """Получить обработчик инструмента"""
        return self.tool_handlers.get(tool_name)
    
    def get_all_tools(self) -> List[Dict]:
        """Получить все инструменты для Mistral"""
        tools = []
        for skill in self.skills.values():
            if skill.enabled:
                for tool in skill.tools:
                    tools.append({
                        "type": "function",
                        "function": tool
                    })
        return tools
    
    def list_skills(self) -> List[Dict]:
        """Список всех навыков"""
        return [
            {
                "name": skill.name,
                "description": skill.description,
                "version": skill.version,
                "enabled": skill.enabled,
                "tools_count": len(skill.tools)
            }
            for skill in self.skills.values()
        ]


# ============================================================
# ВСТРОЕННЫЕ НАВЫКИ
# ============================================================

def create_office_skill() -> tuple:
    """
    Навык для работы с офисными документами
    Создание Word, Excel документов без Google API
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    def create_word_document(title: str, content: str, filename: str = None) -> Dict:
        """Создание Word документа"""
        try:
            doc = Document()
            
            # Заголовок
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Содержимое
            for line in content.split('\n'):
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('| '):
                    # Простая таблица
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    # Добавляем как текст если таблица ещё не начата
                    doc.add_paragraph(' | '.join(cells))
                elif line.strip():
                    doc.add_paragraph(line)
            
            # Сохраняем
            if not filename:
                filename = f"{title.replace(' ', '_')}.docx"
            
            filepath = SKILLS_DIR / "documents" / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(filepath))
            
            return {
                "success": True,
                "filename": filename,
                "filepath": str(filepath),
                "message": f"✅ Документ создан: {filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def create_excel_document(title: str, data: List[List], filename: str = None) -> Dict:
        """Создание Excel документа"""
        try:
            # Проверяем наличие openpyxl
            try:
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment
            except ImportError:
                return {"success": False, "error": "openpyxl not installed. Run: pip install openpyxl"}
            
            wb = Workbook()
            ws = wb.active
            ws.title = title[:31]  # Лимит Excel
            
            for row_idx, row in enumerate(data, 1):
                for col_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    if row_idx == 1:
                        cell.font = Font(bold=True)
            
            if not filename:
                filename = f"{title.replace(' ', '_')}.xlsx"
            
            filepath = SKILLS_DIR / "documents" / filename
            filepath.parent.mkdir(parents=True, exist_ok=True)
            wb.save(str(filepath))
            
            return {
                "success": True,
                "filename": filename,
                "filepath": str(filepath),
                "message": f"✅ Таблица создана: {filename}"
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def read_word_document(filepath: str) -> Dict:
        """Чтение Word документа"""
        try:
            doc = Document(filepath)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return {"success": True, "content": text}
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    skill = Skill(
        name="office",
        description="Создание и редактирование Word и Excel документов",
        version="1.0.0",
        author="HR Bot",
        tools=[
            {
                "name": "create_word_document",
                "description": "Создать Word документ (.docx) с заголовком и содержимым",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string", "description": "Заголовок документа"},
                        "content": {"type": "string", "description": "Содержимое документа (поддерживает Markdown)"},
                        "filename": {"type": "string", "description": "Имя файла (опционально)"}
                    },
                    "required": ["title", "content"]
                }
            },
            {
                "name": "create_excel_document",
                "description": "Создать Excel таблицу (.xlsx) с данными",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string", "description": "Название листа"},
                        "data": {"type": "array", "items": {"type": "array"}, "description": "Данные таблицы (массив массивов)"},
                        "filename": {"type": "string", "description": "Имя файла (опционально)"}
                    },
                    "required": ["title", "data"]
                }
            },
            {
                "name": "read_word_document",
                "description": "Прочитать содержимое Word документа",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "filepath": {"type": "string", "description": "Путь к файлу"}
                    },
                    "required": ["filepath"]
                }
            }
        ],
        dependencies=["python-docx", "openpyxl"]
    )
    
    handlers = {
        "create_word_document": create_word_document,
        "create_excel_document": create_excel_document,
        "read_word_document": read_word_document
    }
    
    return skill, handlers


def create_hr_documents_skill() -> tuple:
    """
    Навык для создания HR документов
    Офферы, welcome-письма, scorecards и т.д.
    """
    from datetime import datetime, timedelta
    
    def create_offer(candidate_name: str, position: str, salary: str, 
                     start_date: str, department: str = "", 
                     company_name: str = "Компания") -> Dict:
        """Создание оффера о приёме на работу"""
        content = f"""# ОФФЕР О ПРИНЯТИИ НА РАБОТУ

**Компания:** {company_name}  
**Дата:** {datetime.now().strftime("%d.%m.%Y")}

---

## Уважаемый(ая) {candidate_name}!

Мы рады предложить Вам должность **{position}** в команде {company_name}.

### Условия предложения:

| Параметр | Значение |
|----------|----------|
| **Должность** | {position} |
| **Отдел** | {department or 'Не указан'} |
| **Тип занятости** | Полная занятость |
| **Дата выхода** | {start_date} |
| **Испытательный срок** | 3 месяца |
| **Зарплата** | {salary} |

### Социальный пакет:
- ДМС после испытательного срока
- Гибкий график
- Оплата обучения

---

Для принятия предложения, пожалуйста, подтвердите в течение 3 рабочих дней.

**С уважением,**  
HR Team  
{company_name}
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Offer_{candidate_name.replace(' ', '_')}.md"
        }
    
    def create_welcome_letter(employee_name: str, position: str, 
                              start_date: str, start_time: str = "10:00",
                              buddy_name: str = "", manager_name: str = "",
                              company_name: str = "Компания") -> Dict:
        """Создание welcome-письма"""
        content = f"""# Добро пожаловать в команду! 🎉

**Привет, {employee_name}!**

Поздравляем с присоединением к команде {company_name}! Мы очень рады, что ты стал(а) частью нашей команды.

---

## 📅 Твой первый день

**Дата:** {start_date}  
**Время:** {start_time}  

---

## 📋 Что взять с собой:
- Паспорт
- ИНН
- СНИЛС
- Диплом об образовании
- Трудовую книжку (если есть)

---

## 🗓️ План первой недели:

**День 1:** Знакомство с командой, оформление документов, настройка рабочего места  
**День 2:** Обучение продуктам и процессам компании  
**День 3:** Знакомство с отделами и ключевыми людьми  
**День 4:** Обучение инструментам и системам  
**День 5:** Постановка первых задач, 1-on-1 с руководителем

---

## 📞 Контакты:

"""
        if buddy_name:
            content += f"- **Buddy:** {buddy_name}\n"
        if manager_name:
            content += f"- **Руководитель:** {manager_name}\n"
        
        content += f"""
---

Если у тебя есть вопросы — не стесняйся писать! Мы всегда на связи.

Добро пожаловать! 🚀
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Welcome_{employee_name.replace(' ', '_')}.md"
        }
    
    def create_rejection_letter(candidate_name: str, position: str,
                                 reason: str = "", keep_in_touch: bool = True,
                                 company_name: str = "Компания") -> Dict:
        """Создание письма с отказом"""
        keep_text = ""
        if keep_in_touch:
            keep_text = "\n- Сохранить Ваше резюме в нашей базе для будущих вакансий\n"
        
        content = f"""# Уважаемый(ая) {candidate_name}!

Благодарим Вас за интерес к вакансии **{position}** в компании {company_name}.

Мы внимательно рассмотрели Вашу кандидатуру и получили большое количество откликов на эту позицию. К сожалению, на данном этапе мы не можем предложить Вам эту должность.

Это решение не означает, что Ваш профессиональный опыт и навыки не представляют ценности — просто в данный момент мы ищем кандидата с другим профилем.

---
{keep_text}
---

Мы желаем Вам успехов в поиске работы и надеемся, что наши пути ещё пересекутся!

С уважением,  
HR Team  
{company_name}
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Rejection_{candidate_name.replace(' ', '_')}.md"
        }
    
    def create_interview_invite(candidate_name: str, position: str,
                                 interview_date: str, interview_time: str,
                                 duration: int = 60, interview_type: str = "онлайн",
                                 company_name: str = "Компания") -> Dict:
        """Создание приглашения на интервью"""
        content = f"""# Приглашение на интервью

**Уважаемый(ая) {candidate_name}!**

Благодарим за интерес к вакансии **{position}** в компании {company_name}.

Мы хотели бы пригласить Вас на интервью.

---

## 📅 Детали интервью:

| Параметр | Значение |
|----------|----------|
| **Дата** | {interview_date} |
| **Время** | {interview_time} |
| **Формат** | {interview_type} |
| **Длительность** | {duration} минут |

---

Пожалуйста, подтвердите своё участие ответным письмом.

Если указанное время неудобно, сообщите нам, и мы подберём другое.

С уважением,  
HR Team  
{company_name}
"""
        return {
            "success": True,
            "content": content,
            "filename": f"Interview_Invite_{candidate_name.replace(' ', '_')}.md"
        }
    
    skill = Skill(
        name="hr_documents",
        description="Создание HR документов: офферы, welcome-письма, отказы, приглашения",
        version="1.0.0",
        author="HR Bot",
        tools=[
            {
                "name": "create_offer_document",
                "description": "Создать оффер о приёме на работу",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "candidate_name": {"type": "string", "description": "Имя кандидата"},
                        "position": {"type": "string", "description": "Должность"},
                        "salary": {"type": "string", "description": "Зарплата"},
                        "start_date": {"type": "string", "description": "Дата выхода"},
                        "department": {"type": "string", "description": "Отдел"},
                        "company_name": {"type": "string", "description": "Название компании"}
                    },
                    "required": ["candidate_name", "position", "salary", "start_date"]
                }
            },
            {
                "name": "create_welcome_letter",
                "description": "Создать welcome-письмо для нового сотрудника",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "employee_name": {"type": "string", "description": "Имя сотрудника"},
                        "position": {"type": "string", "description": "Должность"},
                        "start_date": {"type": "string", "description": "Дата выхода"},
                        "start_time": {"type": "string", "description": "Время выхода"},
                        "buddy_name": {"type": "string", "description": "Имя buddy"},
                        "manager_name": {"type": "string", "description": "Имя руководителя"},
                        "company_name": {"type": "string", "description": "Название компании"}
                    },
                    "required": ["employee_name", "position", "start_date"]
                }
            },
            {
                "name": "create_rejection_letter",
                "description": "Создать письмо с отказом кандидату",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "candidate_name": {"type": "string", "description": "Имя кандидата"},
                        "position": {"type": "string", "description": "Должность"},
                        "reason": {"type": "string", "description": "Причина отказа (опционально)"},
                        "keep_in_touch": {"type": "boolean", "description": "Продолжить общение"},
                        "company_name": {"type": "string", "description": "Название компании"}
                    },
                    "required": ["candidate_name", "position"]
                }
            },
            {
                "name": "create_interview_invite_document",
                "description": "Создать приглашение на интервью",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "candidate_name": {"type": "string", "description": "Имя кандидата"},
                        "position": {"type": "string", "description": "Должность"},
                        "interview_date": {"type": "string", "description": "Дата интервью"},
                        "interview_time": {"type": "string", "description": "Время интервью"},
                        "duration": {"type": "integer", "description": "Длительность в минутах"},
                        "interview_type": {"type": "string", "description": "Тип (онлайн/офис)"},
                        "company_name": {"type": "string", "description": "Название компании"}
                    },
                    "required": ["candidate_name", "position", "interview_date", "interview_time"]
                }
            }
        ]
    )
    
    handlers = {
        "create_offer_document": create_offer,
        "create_welcome_letter": create_welcome_letter,
        "create_rejection_letter": create_rejection_letter,
        "create_interview_invite_document": create_interview_invite
    }
    
    return skill, handlers


def create_web_skill() -> tuple:
    """Навык для работы с вебом"""
    import requests
    
    def fetch_url(url: str) -> Dict:
        """Получить содержимое URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return {
                "success": True,
                "content": response.text[:5000],  # Ограничиваем размер
                "status_code": response.status_code
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def check_url_status(url: str) -> Dict:
        """Проверить статус URL"""
        try:
            response = requests.head(url, timeout=10)
            return {
                "success": True,
                "status_code": response.status_code,
                "accessible": response.status_code < 400
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    skill = Skill(
        name="web",
        description="Работа с веб-страницами: получение содержимого, проверка статуса",
        version="1.0.0",
        author="HR Bot",
        tools=[
            {
                "name": "fetch_url",
                "description": "Получить содержимое веб-страницы по URL",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "URL страницы"}
                    },
                    "required": ["url"]
                }
            },
            {
                "name": "check_url_status",
                "description": "Проверить доступность URL",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "url": {"type": "string", "description": "URL для проверки"}
                    },
                    "required": ["url"]
                }
            }
        ],
        dependencies=["requests"]
    )
    
    handlers = {
        "fetch_url": fetch_url,
        "check_url_status": check_url_status
    }
    
    return skill, handlers


def create_email_skill() -> tuple:
    """Навык для отправки email (через SendGrid или SMTP)"""
    import os
    
    def send_email(to: str, subject: str, body: str, 
                   from_email: str = None, html: bool = False) -> Dict:
        """Отправка email"""
        # Проверяем SendGrid API key
        sendgrid_key = os.getenv("SENDGRID_API_KEY")
        
        if sendgrid_key:
            try:
                import sendgrid
                from sendgrid.helpers.mail import Mail, Email, To, Content
                
                sg = sendgrid.SendGridAPIClient(api_key=sendgrid_key)
                from_email = from_email or os.getenv("SENDGRID_FROM_EMAIL", "noreply@company.com")
                
                message = Mail(
                    from_email=Email(from_email),
                    to_emails=To(to),
                    subject=subject,
                    html_content=Content("text/html" if html else "text/plain", body)
                )
                
                response = sg.send(message)
                return {
                    "success": True,
                    "message": f"✅ Email отправлен на {to}",
                    "status_code": response.status_code
                }
            except Exception as e:
                return {"success": False, "error": f"SendGrid error: {str(e)}"}
        else:
            # Fallback - возвращаем готовое письмо для ручной отправки
            return {
                "success": True,
                "message": "📧 Email готов к отправке (SendGrid не настроен)",
                "email": {
                    "to": to,
                    "subject": subject,
                    "body": body
                },
                "note": "Для автоматической отправки настройте SENDGRID_API_KEY"
            }
    
    skill = Skill(
        name="email",
        description="Отправка email писем через SendGrid или SMTP",
        version="1.0.0",
        author="HR Bot",
        tools=[
            {
                "name": "send_email",
                "description": "Отправить email письмо",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "to": {"type": "string", "description": "Email получателя"},
                        "subject": {"type": "string", "description": "Тема письма"},
                        "body": {"type": "string", "description": "Тело письма"},
                        "from_email": {"type": "string", "description": "Email отправителя"},
                        "html": {"type": "boolean", "description": "HTML формат"}
                    },
                    "required": ["to", "subject", "body"]
                }
            }
        ],
        dependencies=["sendgrid"],
        config={"env_required": ["SENDGRID_API_KEY", "SENDGRID_FROM_EMAIL"]}
    )
    
    handlers = {
        "send_email": send_email
    }
    
    return skill, handlers


# ============================================================
# ИНИЦИАЛИЗАЦИЯ
# ============================================================

# Глобальный реестр навыков
skill_registry = SkillRegistry()


def init_skills():
    """Инициализация всех встроенных навыков"""
    
    # Регистрируем встроенные навыки
    office_skill, office_handlers = create_office_skill()
    skill_registry.register_skill(office_skill, office_handlers)
    
    hr_skill, hr_handlers = create_hr_documents_skill()
    skill_registry.register_skill(hr_skill, hr_handlers)
    
    web_skill, web_handlers = create_web_skill()
    skill_registry.register_skill(web_skill, web_handlers)
    
    email_skill, email_handlers = create_email_skill()
    skill_registry.register_skill(email_skill, email_handlers)
    
    logger.info(f"Initialized {len(skill_registry.skills)} skills with {len(skill_registry.tool_handlers)} tools")
    return skill_registry


# Список доступных навыков для установки
AVAILABLE_SKILLS = {
    "office": {
        "name": "Office Documents",
        "description": "Создание Word и Excel документов",
        "builtin": True
    },
    "hr_documents": {
        "name": "HR Documents",
        "description": "Офферы, welcome-письма, отказы, приглашения",
        "builtin": True
    },
    "web": {
        "name": "Web Tools",
        "description": "Получение содержимого веб-страниц",
        "builtin": True
    },
    "email": {
        "name": "Email",
        "description": "Отправка email через SendGrid",
        "builtin": True,
        "requires": ["SENDGRID_API_KEY"]
    },
    "google_workspace": {
        "name": "Google Workspace",
        "description": "Google Docs, Sheets, Calendar",
        "builtin": False,
        "requires": ["GOOGLE_SERVICE_ACCOUNT_B64"],
        "note": "Уже реализован в google_sheets.py и document_generator.py"
    }
}
